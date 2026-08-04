#!/usr/bin/env python3
"""Generate the REAP Formal Procurement Scorecard training package.

Produces the main PDF/DOCX, one-page quick reference, facilitator notes,
checklist, troubleshooting guide, screenshot index, QA renders and the
client ZIP. Development tooling only - never shipped inside the ZIP.
"""
from __future__ import annotations

import hashlib
import math
import re
import shutil
import textwrap
import zipfile
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Any

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
import fitz
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph, Table, TableStyle

ROOT = Path(__file__).resolve().parents[1]
SHOT = ROOT / "output" / "training-assets" / "screenshots"
BRAND = ROOT / "output" / "training-assets" / "brand" / "reap-solutions-logo.png"
ASSET_OUT = ROOT / "output" / "reap-training-assets"
PDF_DIR = ROOT / "output" / "pdf"
DOC_DIR = ROOT / "output" / "doc"
MD_DIR = ROOT / "output" / "reap-training"
QA_MAIN = ROOT / "output" / "qa" / "main-pages"
QA_QUICK = ROOT / "output" / "qa" / "quick-pages"
ZIP_PATH = ROOT / "output" / "REAP_Scorecard_Training_Package.zip"

MAIN_PDF = PDF_DIR / "REAP_Scorecard_System_Training_Guide.pdf"
MAIN_DOCX = DOC_DIR / "REAP_Scorecard_System_Training_Guide.docx"
QUICK_PDF = PDF_DIR / "REAP_Scorecard_Quick_Reference.pdf"
INDEX_MD = ASSET_OUT / "screenshot-index.md"

PAGE_W, PAGE_H = landscape(A4)  # 841.89 x 595.28
MARGIN = 42.5  # ~15 mm

TEAL = colors.HexColor("#0B5259")
TEAL_DARK = colors.HexColor("#052F32")
IVORY = colors.HexColor("#F7F3EA")
CHARCOAL = colors.HexColor("#161A1F")
MUTED = colors.HexColor("#6B7280")
LINE = colors.HexColor("#D9D2C5")
BRASS = colors.HexColor("#8A5B00")
AMBER_PALE = colors.HexColor("#FBF3E2")
TEAL_PALE = colors.HexColor("#E7F2F2")
WHITE = colors.white
PANEL = colors.HexColor("#FFFDF8")

pdfmetrics.registerFont(TTFont("Disp", "/System/Library/Fonts/Supplemental/Georgia.ttf"))
pdfmetrics.registerFont(TTFont("Disp-Bold", "/System/Library/Fonts/Supplemental/Georgia Bold.ttf"))
pdfmetrics.registerFont(TTFont("Body", "/System/Library/Fonts/Supplemental/Arial.ttf"))
pdfmetrics.registerFont(TTFont("Body-Bold", "/System/Library/Fonts/Supplemental/Arial Bold.ttf"))

styles = {
    "body": ParagraphStyle("body", fontName="Body", fontSize=9.2, leading=12.4, textColor=CHARCOAL),
    "body_sm": ParagraphStyle("body_sm", fontName="Body", fontSize=8.4, leading=11.2, textColor=CHARCOAL),
    "muted": ParagraphStyle("muted", fontName="Body", fontSize=8.2, leading=11, textColor=MUTED),
    "step": ParagraphStyle("step", fontName="Body", fontSize=9.2, leading=12.4, textColor=CHARCOAL),
    "caption": ParagraphStyle("caption", fontName="Body", fontSize=7.8, leading=10, textColor=MUTED),
    "panel": ParagraphStyle("panel", fontName="Body", fontSize=8.2, leading=11, textColor=CHARCOAL),
    "panel_title": ParagraphStyle("panel_title", fontName="Body-Bold", fontSize=8.0, leading=10, textColor=TEAL_DARK),
    "h_disp": ParagraphStyle("h_disp", fontName="Disp-Bold", fontSize=22, leading=26, textColor=TEAL_DARK),
}


def mm(n: float) -> float:
    return n * 72 / 25.4


def shot_path(name: str) -> Path | None:
    if not name:
        return None
    p = SHOT / name
    if not p.suffix:
        p = SHOT / f"{name}.png"
    return p if p.exists() else None


@dataclass
class Page:
    code: str
    section: str
    title: str
    purpose: str = ""
    kind: str = "ops"  # cover|contents|concept|journey|divider|ops|table|checklist|close
    steps: list[str] = field(default_factory=list)
    checks: list[str] = field(default_factory=list)
    mistake: str = ""
    expected: str = ""
    images: list[str] = field(default_factory=list)
    captions: list[str] = field(default_factory=list)
    fictional: bool = False
    body: list[str] = field(default_factory=list)
    bullets: list[str] = field(default_factory=list)
    table: list[list[str]] = field(default_factory=list)
    note: str = ""


def build_pages() -> list[Page]:
    P: list[Page] = []
    # ----- SECTION 1 -----
    P.append(Page("01", "REAP SOLUTIONS", "REAP Formal Procurement Scorecard",
                  purpose="System training and operations guide", kind="cover"))
    P.append(Page("02", "CONTENTS", "Contents", kind="contents"))
    P.append(Page("03", "SECTION 1 · INTRODUCTION", "About this guide", kind="concept",
                  purpose="How to use this document during training and afterward as the operating manual.",
                  body=[
                      "This guide is the definitive training and operations reference for the REAP Formal Procurement Scorecard.",
                      "It is written for REAP consultants, trainers and operators who create company records, run procurement assessments, interpret scores and issue client reports.",
                      "Use it live during instructor-led sessions. Keep it beside the system afterward as the permanent operating manual and quick-reference companion.",
                  ],
                  bullets=[
                      "Who should use it — REAP practitioners, trainers, and authorised client operators.",
                      "During training — follow the numbered steps while the facilitator drives the demonstration.",
                      "After training — reopen any section as a task card when you next run an assessment.",
                      "Source of truth — the values shown in the live system always take precedence over worked examples.",
                  ]))
    P.append(Page("04", "SECTION 1 · INTRODUCTION", "Platform overview", kind="ops",
                  purpose="The Formal Procurement Scorecard turns supplier-spend evidence into a defendable preferential-procurement result.",
                  steps=[
                      "Open the public REAP Scorecard page to understand the product purpose.",
                      "Sign in with an authorised account.",
                      "Create or open the company that owns the assessment.",
                      "Run a procurement assessment for the correct reporting year.",
                      "Confirm TMPS, review the live score, save, then export the client report.",
                  ],
                  checks=["You are in the Formal Procurement Scorecard, not any other module.",
                          "The company and assessment year are known before data is uploaded."],
                  mistake="Do not start an assessment without a company record — assessments are always linked to a company.",
                  expected="You can name the full journey from company creation through PDF download.",
                  images=["01-public-page.png"],
                  captions=["Public REAP Scorecard page — the formal product entry point."],
                  fictional=False))
    P.append(Page("05", "SECTION 1 · INTRODUCTION", "Key system concepts", kind="concept",
                  purpose="Shared vocabulary used throughout the Formal Procurement Scorecard.",
                  body=[
                      "Company record — the organisation profile that owns assessments, contacts and history.",
                      "Assessment — a year-specific procurement calculation saved against a company.",
                      "Supplier register — the list of suppliers, spend amounts, B-BBEE levels and ownership flags used in the assessment.",
                      "B-BBEE level — the supplier’s status level (1–8 or Non-compliant). Recognition percentage is derived from this level.",
                      "Recognised spend — supplier spend multiplied by the recognition percentage for that level.",
                      "TMPS — Total Measured Procurement Spend: the denominator that converts recognised spend into category percentages.",
                      "Procurement score — the sum of capped category points, out of a maximum of 29.",
                      "Category target / maximum points — the percentage target and point ceiling for each preferential-procurement category.",
                      "Client report — the printable and downloadable record of the saved assessment.",
                  ],
                  note="The procurement result is not the company’s complete B-BBEE score. It covers preferential procurement only."))
    P.append(Page("06", "SECTION 1 · INTRODUCTION", "Worked example: recognised spend", kind="concept",
                  purpose="How supplier spend and recognition percentage combine. Figures are fictional.",
                  body=[
                      "Example (fictional): a Level 2 supplier with B-BBEE spend of R 1 000 000.",
                      "Recognition for Level 2 is 125%. Recognised spend = R 1 000 000 × 1.25 = R 1 250 000.",
                      "If that supplier is also 51% Black Owned, the same recognised amount contributes to the Black Owned category.",
                      "Category achieved % = recognised category spend ÷ TMPS. Points = min(achieved% ÷ target% × max points, max points).",
                  ],
                  table=[
                      ["Level", "Recognition", "Example spend", "Recognised spend"],
                      ["Level 1", "135%", "R 1 000 000", "R 1 350 000"],
                      ["Level 2", "125%", "R 1 000 000", "R 1 250 000"],
                      ["Level 4", "100%", "R 1 000 000", "R 1 000 000"],
                      ["Level 8", "10%", "R 1 000 000", "R 100 000"],
                      ["Non-compliant", "0%", "R 1 000 000", "R 0"],
                  ],
                  note="Illustrative recognition percentages are not legal advice. Evidence must support supplier levels and ownership. The system’s displayed values remain the source of truth."))
    P.append(Page("07", "SECTION 1 · INTRODUCTION", "Full assessment journey", kind="journey",
                  purpose="Every formal assessment follows the same operating path.",
                  bullets=[
                      "Create or open company",
                      "Start assessment",
                      "Select assessment year",
                      "Add supplier data",
                      "Map columns",
                      "Validate and review",
                      "Confirm TMPS",
                      "Review score",
                      "Save assessment",
                      "Generate report",
                      "Download PDF",
                      "Review activity",
                  ]))

    # ----- SECTION 2 -----
    P.append(Page("08", "SECTION 2", "Access and navigation", kind="divider",
                  purpose="Sign in securely, orient yourself, and leave the session cleanly."))
    P.append(Page("09", "SECTION 2 · ACCESS", "Open the public Scorecard page", kind="ops",
                  purpose="Confirm you are entering the Formal Procurement Scorecard product.",
                  steps=["Open the approved REAP Scorecard web address.",
                         "Read the product purpose on the public page.",
                         "Select Sign in to Scorecard when ready."],
                  checks=["The browser address is the approved REAP site.",
                          "No developer tools or unrelated browser tabs are visible during training."],
                  mistake="Do not bookmark an outdated staging URL for client work.",
                  expected="The public page presents the Formal Procurement Scorecard value proposition.",
                  images=["01-public-page.png"],
                  captions=["Public product page."]))
    P.append(Page("10", "SECTION 2 · ACCESS", "Create an account", kind="ops",
                  purpose="New operators register once, then sign in with email and password.",
                  steps=["From the sign-in screen, select Create account.",
                         "Enter full name, work email and a strong password.",
                         "Confirm the password and select Create account.",
                         "Return to sign in if the organisation already has an approved user."],
                  checks=["The email belongs to an authorised operator.",
                          "Password requirements shown on screen are met."],
                  mistake="Do not create a second account when a password reset is all that is needed.",
                  expected="The account is ready for sign-in.",
                  images=["02-registration.png", "02b-registration-detail.png"],
                  captions=["Create your account.", "Required registration fields."],
                  fictional=True))
    P.append(Page("11", "SECTION 2 · ACCESS", "Sign in", kind="ops",
                  purpose="Secure sign-in protects company records, assessments and reports.",
                  steps=["Open the sign-in screen.",
                         "Enter the authorised email and password.",
                         "Select Sign in with email.",
                         "Confirm the Dashboard opens before starting work."],
                  checks=["The correct account is used.",
                          "No password appears in screenshots or training notes."],
                  mistake="Do not share passwords in chat, slides or printed handouts.",
                  expected="The Dashboard loads for the signed-in operator.",
                  images=["03-login.png", "03b-login-form-detail.png"],
                  captions=["Sign-in screen.", "Email and password fields."],
                  fictional=True))
    P.append(Page("12", "SECTION 2 · ACCESS", "Complete first-time onboarding", kind="ops",
                  purpose="New accounts see a setup checklist and an optional guided walkthrough.",
                  steps=["Read the Complete your setup checklist on first sign-in.",
                         "Follow Add your first company when the workspace is empty.",
                         "Open Need help? if you want the interactive guided walkthrough.",
                         "Skip or finish the guide when you are oriented."],
                  checks=["The checklist matches your current progress.",
                          "The guide does not change saved data."],
                  mistake="Do not ignore the company step — assessments require a company record.",
                  expected="You know where Companies, New Procurement Assessment and Need help? live.",
                  images=["04-onboarding.png", "04b-guided-walkthrough.png", "04c-guided-walkthrough-detail.png"],
                  captions=["First-time setup checklist.", "Guided walkthrough overlay.", "Walkthrough step card."],
                  fictional=True))
    P.append(Page("13", "SECTION 2 · ACCESS", "Use the Dashboard and main navigation", kind="ops",
                  purpose="The Dashboard is the home for portfolio status, recent assessments and navigation.",
                  steps=["Open Dashboard from the left navigation.",
                         "Use Companies, Activity, New Company and New Procurement Assessment as needed.",
                         "Open Settings for profile and Help Centre.",
                         "Review recent assessments only after confirming company and year."],
                  checks=["The navigation labels match this guide.",
                          "Recent work belongs to the intended company."],
                  mistake="Do not assume the newest assessment is the one required for today’s task.",
                  expected="You can reach every formal workflow area from the Dashboard.",
                  images=["05-dashboard.png", "05b-main-navigation-detail.png", "05c-dashboard-portfolio.png"],
                  captions=["Dashboard with saved work.", "Main navigation.", "Portfolio metrics."],
                  fictional=True))
    P.append(Page("14", "SECTION 2 · ACCESS", "Sign out securely", kind="ops",
                  purpose="End the session when work is complete, especially on shared devices.",
                  steps=["Open the user menu or the sidebar sign-out control.",
                         "Select Sign out.",
                         "Confirm the sign-in screen returns."],
                  checks=["No assessment edit screen remains open on a shared computer."],
                  mistake="Do not leave a signed-in session unattended during a break.",
                  expected="The session ends and the sign-in screen is shown.",
                  images=["26-sign-out-detail.png"],
                  captions=["Sign out control."],
                  fictional=True))

    # ----- SECTION 3 -----
    P.append(Page("15", "SECTION 3", "Company management", kind="divider",
                  purpose="Every assessment is owned by a company record. Keep company data accurate and unique."))
    P.append(Page("16", "SECTION 3 · COMPANIES", "Open the company directory", kind="ops",
                  purpose="Company records keep procurement assessments attached to the correct organisation.",
                  steps=["Select Companies in the left navigation.",
                         "Review each organisation name, industry and contact.",
                         "Open View profile on the intended company.",
                         "If the directory is empty, create a company first."],
                  checks=["The legal or trading name is exact.",
                          "You are not opening a similarly named duplicate."],
                  mistake="The user directory does not provide a search box — read each record carefully before opening it.",
                  expected="The correct company profile opens.",
                  images=["06-company-directory-empty.png", "06b-company-directory.png", "06c-company-record-detail.png"],
                  captions=["Empty directory.", "Directory with a company.", "Company record detail."],
                  fictional=True))
    P.append(Page("17", "SECTION 3 · COMPANIES", "Create a company", kind="ops",
                  purpose="Capture the organisation once so all assessments hang off a single profile.",
                  steps=["Select New Company.",
                         "Enter Company name, Contact name, Work email and Phone.",
                         "Add Industry and Notes when useful.",
                         "Select Save company."],
                  checks=["Required fields are complete.",
                          "The name matches the client’s legal or trading name."],
                  mistake="Do not create a second company for a spelling variation of an existing client.",
                  expected="The company profile opens with the saved details.",
                  images=["07-create-company.png", "07b-create-company-completed.png", "07c-required-fields-detail.png"],
                  captions=["New company form.", "Completed demonstration details.", "Required fields."],
                  fictional=True))
    P.append(Page("18", "SECTION 3 · COMPANIES", "Review the company profile and history", kind="ops",
                  purpose="The profile is the home for company details and saved procurement assessments.",
                  steps=["Confirm industry, contact, email and phone.",
                         "Use Edit company when details change.",
                         "Start New Procurement Assessment from the profile when ready.",
                         "Open an existing assessment from the history list to review or edit."],
                  checks=["Contact details are current.",
                          "Assessment history shows the expected years."],
                  mistake="Do not start a new assessment for the wrong company profile.",
                  expected="Profile details and assessment history are trustworthy.",
                  images=["08-company-profile.png", "08b-company-details-detail.png", "08d-company-assessment-history.png"],
                  captions=["Company profile.", "Stored company details.", "Assessment history on the profile."],
                  fictional=True))

    # ----- SECTION 4 -----
    P.append(Page("19", "SECTION 4", "Procurement assessments", kind="divider",
                  purpose="Capture supplier evidence, confirm TMPS, review the live score and save a defendable assessment."))
    P.append(Page("20", "SECTION 4 · ASSESSMENTS", "Start a new assessment", kind="ops",
                  purpose="Assessments are created from a company and locked to a reporting year.",
                  steps=["Open the company profile.",
                         "Select New Procurement Assessment.",
                         "Confirm the company name in the assessment header.",
                         "Enter the assessment year (2000–2100)."],
                  checks=["The company name is correct.",
                          "The year matches the reporting period of the supplier workbook."],
                  mistake="Do not reuse last year’s assessment when a new reporting year is required — create a new assessment.",
                  expected="The procurement input form opens for the correct company and year.",
                  images=["09-new-assessment.png", "09b-assessment-year-detail.png", "09c-how-this-works-detail.png"],
                  captions=["New assessment workspace.", "Assessment year field.", "How this works panel."],
                  fictional=True))
    P.append(Page("21", "SECTION 4 · ASSESSMENTS", "Confirm the TMPS method", kind="ops",
                  purpose="TMPS is the denominator for every category percentage. Choose the method that matches the evidence.",
                  steps=["Open TMPS (measured procurement).",
                         "Select Calculated TMPS to use the inclusions and exclusions pad, or Use supplier spend as TMPS when that is the agreed method.",
                         "Enter inclusion and exclusion amounts for Calculated TMPS.",
                         "Confirm the live scoring denominator is greater than zero before saving."],
                  checks=["The method matches the client’s TMPS schedule.",
                          "Calculated TMPS is not left at R0 when that method is selected."],
                  mistake="Do not save while the scoring denominator is zero.",
                  expected="A positive TMPS denominator is shown in the live summary.",
                  images=["15-tmps-method.png", "15b-tmps-method-detail.png", "15d-tmps-inclusions-detail.png", "15f-calculated-tmps-detail.png"],
                  captions=["TMPS method selection.", "Denominator choices.", "Inclusions pad.", "Calculated TMPS total."],
                  fictional=True,
                  note="Why TMPS matters: without a correct denominator, category percentages and points are not defendable."))
    P.append(Page("22", "SECTION 4 · ASSESSMENTS", "Upload the supplier workbook", kind="ops",
                  purpose="Import supplier rows from an Excel workbook (.xlsx or .xls).",
                  steps=["Prepare a supplier register sheet with names and spend amounts.",
                         "In Supplier capture, open Upload a procurement workbook.",
                         "Drop or browse to the file.",
                         "Confirm the detected sheet used for suppliers."],
                  checks=["The workbook contains a supplier-level register.",
                          "Finance-only TMPS sheets are not selected as the supplier sheet."],
                  mistake="Do not upload a summary sheet that has no supplier rows.",
                  expected="Detected file summary shows the workbook, sheet, and rows read.",
                  images=["10-upload-workbook.png", "10b-upload-panel-detail.png", "10c-uploaded-file-summary.png", "10d-sheet-selection-detail.png"],
                  captions=["Upload panel.", "Drop zone.", "Detected file summary.", "Sheet used for suppliers."],
                  fictional=True))
    P.append(Page("23", "SECTION 4 · ASSESSMENTS", "Map workbook columns", kind="ops",
                  purpose="Required fields must point at the correct columns before suppliers can be applied.",
                  steps=["Open Column mapping.",
                         "Confirm Supplier name and Spend amount are Found.",
                         "Map optional fields for level, ownership, designated groups and supplier type when present.",
                         "Use Find column name to narrow long header lists."],
                  checks=["Required mappings show Found.",
                          "Optional mappings improve category allocation when evidence exists."],
                  mistake="Do not leave a required field on Needs mapping.",
                  expected="Required columns are mapped and optional fields are set where available.",
                  images=["11-column-mapping.png", "11b-column-mapping-fields-detail.png", "11c-mapping-status-detail.png"],
                  captions=["Column mapping workspace.", "Field mapping table.", "Status for a required field."],
                  fictional=True))
    P.append(Page("24", "SECTION 4 · ASSESSMENTS", "Validate imported rows", kind="ops",
                  purpose="Review loaded suppliers, warnings and skipped rows before applying the import.",
                  steps=["Read the Procurement upload result summary.",
                         "Open Warnings and notes for skipped or invalid rows.",
                         "Correct the workbook or plan manual edits for rejected rows.",
                         "Select Apply suppliers to this assessment."],
                  checks=["Supplier count matches expectations.",
                          "Skipped rows are understood before continuing."],
                  mistake="Do not ignore zero-spend or non-numeric spend warnings — those rows will not score.",
                  expected="Suppliers appear in the supplier register.",
                  images=["12-validation-summary.png", "12b-ready-row-count-detail.png", "12c-import-warnings.png", "12d-skipped-rows-detail.png"],
                  captions=["Import validation summary.", "Suppliers loaded.", "Warnings panel.", "Skipped row detail."],
                  fictional=True))
    P.append(Page("25", "SECTION 4 · ASSESSMENTS", "Paste supplier data when needed", kind="ops",
                  purpose="Paste is available for quick TSV/CSV entry when a full workbook is not required.",
                  steps=["Open Paste from spreadsheet.",
                         "Paste rows with supplier name and B-BBEE spend in the first two columns.",
                         "Select Import pasted rows.",
                         "Review import notes for skipped lines."],
                  checks=["Header rows are skipped automatically.",
                          "Each imported line has a positive spend amount."],
                  mistake="Do not paste category totals instead of supplier lines.",
                  expected="Pasted suppliers appear in the register.",
                  images=["10f-paste-import.png", "10g-paste-import-detail.png"],
                  captions=["Paste import panel.", "Paste area and accepted columns."],
                  fictional=True))
    P.append(Page("26", "SECTION 4 · ASSESSMENTS", "Review and edit the supplier register", kind="ops",
                  purpose="Confirm levels, ownership and spend before trusting the live score.",
                  steps=["Open Suppliers and B-BBEE spend.",
                         "Use Find a supplier to filter the list.",
                         "Expand a row to edit details, classification and ownership flags.",
                         "Add a supplier row manually when a record is missing."],
                  checks=["B-BBEE level matches supporting evidence.",
                          "Ownership flags are only set when evidenced."],
                  mistake="Do not mark Black Owned or Designated Group without evidence.",
                  expected="The register reflects the agreed supplier evidence pack.",
                  images=["13-supplier-register.png", "13c-supplier-search.png", "14-edit-supplier.png", "14c-level-recognition-detail.png", "14d-ownership-detail.png"],
                  captions=["Supplier register.", "Filtered search.", "Expanded supplier editor.", "Level and recognition.", "Ownership flags."],
                  fictional=True))
    P.append(Page("27", "SECTION 4 · ASSESSMENTS", "Review the live score preview and save", kind="ops",
                  purpose="The live preview shows estimated points before the assessment is saved.",
                  steps=["Review Live summary tiles for denominator, recognised spend and estimated points.",
                         "Open Procurement Score Preview and check each category.",
                         "Confirm TMPS, suppliers and preview one last time.",
                         "Select Save procurement assessment."],
                  checks=["Denominator is greater than zero.",
                          "At least one supplier with positive spend exists."],
                  mistake="Do not save while estimated points still look wrong — fix suppliers or TMPS first.",
                  expected="The saved assessment page opens with the calculated result.",
                  images=["16c-live-summary-detail.png", "16-live-score-preview.png", "16b-score-preview-detail.png", "16e-save-action-detail.png"],
                  captions=["Live summary.", "Score preview.", "Category table.", "Save action."],
                  fictional=True))
    P.append(Page("28", "SECTION 4 · ASSESSMENTS", "Reopen and recalculate an assessment", kind="ops",
                  purpose="Edit an assessment when evidence changes, then recalculate.",
                  steps=["Open the assessment from the company profile or Dashboard.",
                         "Select Edit.",
                         "Adjust suppliers or TMPS as required.",
                         "Select Save changes & recalculate."],
                  checks=["You are editing the correct year.",
                          "Recalculation completes without a zero denominator."],
                  mistake="Do not edit a prior year when the client asked for the current year.",
                  expected="Updated points appear on the assessment page.",
                  images=["22-edit-assessment.png", "22b-recalculate-detail.png"],
                  captions=["Edit assessment workspace.", "Save changes & recalculate."],
                  fictional=True))

    # ----- SECTION 5 -----
    P.append(Page("29", "SECTION 5", "Results and reporting", kind="divider",
                  purpose="Interpret the saved result, explain category gaps and issue the client report."))
    P.append(Page("30", "SECTION 5 · RESULTS", "Read the assessment summary", kind="ops",
                  purpose="The summary states total score, TMPS, recognised spend and procurement level.",
                  steps=["Open the saved assessment.",
                         "Read total score out of 29 and the procurement level.",
                         "Confirm measured procurement spend and recognised B-BBEE spend.",
                         "If a prior assessment exists, review the year-on-year comparison."],
                  checks=["Company and year in the header are correct.",
                          "Comparison baseline year is understood before discussing movement."],
                  mistake="Do not present the procurement level as the entity’s overall B-BBEE level.",
                  expected="You can explain score, TMPS and recognised spend in one sentence each.",
                  images=["17-saved-assessment.png", "17b-assessment-summary-detail.png", "17c-year-on-year.png"],
                  captions=["Saved assessment.", "Assessment summary.", "Compared to previous assessment."],
                  fictional=True))
    P.append(Page("31", "SECTION 5 · RESULTS", "Executive scorecard and supplier breakdown", kind="ops",
                  purpose="Executive metrics and the recognised supplier breakdown support client conversations.",
                  steps=["Open Executive scorecard.",
                         "Review procurement score, level, TMPS and recognised spend.",
                         "Open the Preferential Procurement table for category points.",
                         "Review Recognised supplier breakdown for contribution by supplier."],
                  checks=["Points total matches the summary.",
                          "Supplier recognition percentages align with levels."],
                  mistake="Do not hide non-compliant or low-recognition suppliers when explaining gaps.",
                  expected="You can point to the suppliers driving recognised spend.",
                  images=["18-executive-scorecard.png", "18d-scorecard-table-detail.png", "18e-supplier-breakdown.png", "18f-supplier-breakdown-detail.png"],
                  captions=["Executive scorecard.", "Preferential procurement table.", "Supplier breakdown.", "Recognised spend columns."],
                  fictional=True))
    P.append(Page("32", "SECTION 5 · RESULTS", "Category performance and recommendations", kind="ops",
                  purpose="Category targets, gaps and recommended actions explain where to improve.",
                  steps=["Open Category performance.",
                         "Note Strong, Near target and Action required statuses.",
                         "Open Detailed category breakdown for recognised value against TMPS.",
                         "Read Recommended improvement actions with the client."],
                  checks=["Targets and maximum points are understood before discussing gaps.",
                          "Recommendations are tied to evidenced categories."],
                  mistake="Do not promise points that the current supplier evidence cannot support.",
                  expected="Each weak category has a clear next action.",
                  images=["19-category-performance.png", "19b-category-performance-detail.png", "19e-detailed-category-analysis.png", "19g-recommendations.png"],
                  captions=["Category performance.", "Targets and points.", "Detailed analysis.", "Recommended actions."],
                  fictional=True,
                  table=[
                      ["Category", "Target", "Max points"],
                      ["All B-BBEE Suppliers", "80%", "5"],
                      ["All QSEs", "15%", "3"],
                      ["All EMEs", "15%", "4"],
                      ["51% Black Owned", "50%", "11"],
                      ["30% Black Women Owned", "12%", "4"],
                      ["51% Black Designated Groups", "2%", "2"],
                      ["Total", "—", "29"],
                  ]))
    P.append(Page("33", "SECTION 5 · RESULTS", "Generate the client report and PDF", kind="ops",
                  purpose="Issue a printable client report and download the PDF.",
                  steps=["Select View report.",
                         "Review the report content on screen.",
                         "Select Download PDF.",
                         "Confirm the file saves with a professional REAP procurement scorecard filename."],
                  checks=["The report matches the saved assessment.",
                          "The PDF opens and is readable."],
                  mistake="Do not send a screenshot of the screen when the PDF export is available.",
                  expected="A PDF named for the company and assessment is downloaded.",
                  images=["20-client-report.png", "20b-report-toolbar-detail.png", "21-pdf-download.png", "21b-pdf-download-complete.png"],
                  captions=["Client report.", "Report toolbar.", "PDF download in progress.", "Downloaded report page."],
                  fictional=True))
    P.append(Page("34", "SECTION 5 · RESULTS", "Review activity", kind="ops",
                  purpose="Activity shows recent create and update actions for accountability.",
                  steps=["Open Activity from the left navigation.",
                         "Confirm the company and assessment events appear.",
                         "Use activity when reconstructing what changed in a session."],
                  checks=["Timestamps match the training or operating session."],
                  mistake="Do not treat activity as a substitute for the assessment report.",
                  expected="Key create and update events are visible.",
                  images=["23-activity.png", "23b-activity-detail.png"],
                  captions=["Activity page.", "Recent activity detail."],
                  fictional=True))

    # ----- SECTION 6 -----
    P.append(Page("35", "SECTION 6", "Operations and support", kind="divider",
                  purpose="Keep the workspace configured, follow the operating checklists, and escalate cleanly."))
    P.append(Page("36", "SECTION 6 · OPERATIONS", "Profile settings and Help Centre", kind="ops",
                  purpose="Maintain the operator profile and reopen interactive guides when needed.",
                  steps=["Open Settings → Profile to update name and photo.",
                         "Open Help Centre for interactive guides.",
                         "Launch Procurement scorecard guide or Workbook upload guide during training refreshers.",
                         "Return to the Dashboard when finished."],
                  checks=["Profile name is professional for client-facing sessions."],
                  mistake="Do not disable help before a first-time user has completed onboarding.",
                  expected="Help guides are available without changing saved assessments.",
                  images=["24-settings.png", "25-help-centre.png", "25b-help-guides-detail.png"],
                  captions=["Profile settings.", "Help Centre.", "Interactive guides."],
                  fictional=True))
    P.append(Page("37", "SECTION 6 · OPERATIONS", "Operating checklists", kind="checklist",
                  purpose="Use these checks before assessment work, before saving, and before export.",
                  bullets=[
                      "Before assessment — company selected; year agreed; workbook ready; TMPS schedule available; evidence pack available.",
                      "Before saving — TMPS denominator > 0; required columns mapped; skipped rows reviewed; ownership flags evidenced; live preview reviewed.",
                      "Before exporting — saved assessment opened; company and year confirmed; PDF download tested; filename professional; confidential data excluded from training copies.",
                      "Session closeout — report shared if required; activity spot-checked; signed out on shared devices; open questions logged.",
                  ]))
    P.append(Page("38", "SECTION 6 · OPERATIONS", "Troubleshooting essentials", kind="ops",
                  purpose="Resolve common issues before escalating.",
                  steps=["Cannot sign in — confirm email, use Forgot password, verify account exists.",
                         "Workbook rejected — use .xlsx/.xls, keep under the server size limit, ensure a supplier sheet exists.",
                         "Rows skipped — fix non-numeric or zero spend, remapping if needed.",
                         "Unexpected score — recheck TMPS method, levels, ownership flags and recognised spend.",
                         "PDF not downloading — stay on the report page, retry Download PDF, check browser download settings."],
                  checks=["Exact error text is captured before contacting support."],
                  mistake="Do not recreate the company as a workaround for a mapping problem.",
                  expected="Most issues are resolved with the checks above; unresolved items include the support information list.",
                  images=["12c-import-warnings.png"],
                  captions=["Import warnings are the first place to look for data problems."],
                  fictional=True))
    P.append(Page("39", "SECTION 6 · OPERATIONS", "Information required when requesting support", kind="concept",
                  purpose="Give support enough context to reproduce the issue quickly.",
                  bullets=[
                      "Company name",
                      "Assessment year",
                      "Workbook filename",
                      "Screenshot of the screen showing the problem",
                      "Exact error message",
                      "Steps completed before the error",
                  ],
                  body=["Do not include passwords, access tokens or confidential client workbooks in the support pack unless a secure channel has been agreed."]))
    P.append(Page("40", "SECTION 6 · OPERATIONS", "Closing page", kind="close",
                  purpose="Operate with evidence, confirm TMPS, and issue only what the Formal Procurement Scorecard can defend.",
                  body=[
                      "The REAP Formal Procurement Scorecard exists to turn supplier-spend evidence into a clear preferential-procurement result.",
                      "Create accurate company records. Map columns carefully. Confirm TMPS. Review the live score. Save. Export the report. Check activity.",
                      "When in doubt, return to this guide’s operating pages and the one-page quick reference.",
                  ]))
    return P


PROHIBITED = re.compile(
    r"aberdare|live procurement|scenario simulat|actual.?versus.?projected|actual vs projected|"
    r"client workspace|code 400|sap modell?ing|procurement control|monthly sap",
    re.I,
)


class GuidePDF:
    def __init__(self, pages: list[Page], path: Path):
        self.pages = pages
        self.path = path
        self.c = canvas.Canvas(str(path), pagesize=landscape(A4))
        self.c.setTitle("REAP Formal Procurement Scorecard — System Training Guide")
        self.c.setAuthor("REAP Solutions")
        self.c.setSubject("Formal Procurement Scorecard training and operations")
        self.page_no = 0
        self.used_images: dict[str, list[int]] = {}

    def draw(self):
        for p in self.pages:
            self.page_no += 1
            getattr(self, f"_kind_{p.kind}", self._kind_ops)(p)
            self._footer(p)
            self.c.showPage()
        self.c.save()

    def _bg(self, ivory=True):
        self.c.setFillColor(IVORY if ivory else WHITE)
        self.c.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)

    def _footer(self, p: Page):
        y = MARGIN - 18
        self.c.setStrokeColor(LINE)
        self.c.setLineWidth(0.6)
        self.c.line(MARGIN, MARGIN - 8, PAGE_W - MARGIN, MARGIN - 8)
        self.c.setFillColor(MUTED)
        self.c.setFont("Body", 7.5)
        self.c.drawString(MARGIN, y, "REAP Formal Procurement Scorecard · Operations Ledger")
        self.c.drawRightString(PAGE_W - MARGIN, y, f"{p.code}  ·  {self.page_no}")

    def _section_label(self, text: str, y: float):
        self.c.setFillColor(TEAL)
        self.c.setFont("Body-Bold", 8)
        self.c.drawString(MARGIN, y, text.upper())
        self.c.setStrokeColor(TEAL)
        self.c.setLineWidth(1)
        self.c.line(MARGIN, y - 4, MARGIN + 54, y - 4)

    def _title(self, text: str, y: float, size=22):
        self.c.setFillColor(TEAL_DARK)
        self.c.setFont("Disp-Bold", size)
        for line in textwrap.wrap(text, width=48 if size >= 20 else 60):
            self.c.drawString(MARGIN, y, line)
            y -= size + 4
        return y

    def _purpose(self, text: str, y: float):
        if not text:
            return y
        p = Paragraph(text, styles["muted"])
        w, h = p.wrap(PAGE_W - 2 * MARGIN, 60)
        p.drawOn(self.c, MARGIN, y - h)
        return y - h - 10

    def _rule(self, y: float):
        self.c.setStrokeColor(LINE)
        self.c.setLineWidth(0.7)
        self.c.line(MARGIN, y, PAGE_W - MARGIN, y)
        return y - 12

    def _steps_table(self, steps: list[str], x: float, y: float, width: float) -> float:
        if not steps:
            return y
        data = []
        for i, step in enumerate(steps, 1):
            num = Paragraph(f"<b>{i:02d}</b>", ParagraphStyle("n", fontName="Body-Bold", fontSize=9, textColor=TEAL, leading=12))
            txt = Paragraph(step, styles["step"])
            data.append([num, txt])
        t = Table(data, colWidths=[28, width - 28])
        t.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 2),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LINEBELOW", (0, 0), (-1, -2), 0.4, LINE),
        ]))
        w, h = t.wrap(width, y - MARGIN - 20)
        t.drawOn(self.c, x, y - h)
        return y - h - 8

    def _panel(self, title: str, body: str, x: float, y: float, w: float, fill=TEAL_PALE, title_color=TEAL_DARK) -> float:
        p_title = Paragraph(title.upper(), ParagraphStyle("pt", fontName="Body-Bold", fontSize=7.6, textColor=title_color, leading=10))
        p_body = Paragraph(body, styles["panel"])
        tw, th = p_title.wrap(w - 16, 40)
        bw, bh = p_body.wrap(w - 16, 200)
        height = th + bh + 18
        self.c.setFillColor(fill)
        self.c.setStrokeColor(LINE)
        self.c.roundRect(x, y - height, w, height, 4, fill=1, stroke=1)
        p_title.drawOn(self.c, x + 8, y - 10 - th)
        p_body.drawOn(self.c, x + 8, y - 14 - th - bh)
        return y - height - 8

    def _draw_image(self, name: str, x: float, y_top: float, max_w: float, max_h: float, caption: str = "", fictional: bool = False) -> float:
        path = shot_path(name)
        if not path:
            self.c.setFillColor(MUTED)
            self.c.setFont("Body", 8)
            self.c.drawString(x, y_top - 12, f"[Screenshot unavailable: {name}]")
            return y_top - 24
        y = y_top
        if fictional:
            self.c.setFillColor(BRASS)
            self.c.setFont("Body-Bold", 7.2)
            self.c.drawString(x, y, "FICTIONAL DEMONSTRATION DATA")
            y -= 12
        im = Image.open(path)
        iw, ih = im.size
        scale = min(max_w / iw, max_h / ih, 1.0)
        dw, dh = iw * scale, ih * scale
        # Prefer not to upscale; reportlab draws in points
        self.c.setStrokeColor(LINE)
        self.c.setLineWidth(0.6)
        self.c.setFillColor(WHITE)
        self.c.rect(x - 2, y - dh - 2, dw + 4, dh + 4, fill=1, stroke=1)
        self.c.drawImage(ImageReader(str(path)), x, y - dh, width=dw, height=dh, preserveAspectRatio=True, mask="auto")
        self.used_images.setdefault(name, []).append(self.page_no)
        y = y - dh - 4
        if caption:
            cap = Paragraph(f"Figure — {caption}", styles["caption"])
            cw, ch = cap.wrap(max_w, 40)
            cap.drawOn(self.c, x, y - ch)
            y -= ch + 6
        # effective DPI note for preflight: (iw / (dw/72))
        return y

    def _kind_cover(self, p: Page):
        self.c.setFillColor(TEAL_DARK)
        self.c.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)
        self.c.setFillColor(IVORY)
        self.c.rect(0, 0, mm(18), PAGE_H, fill=1, stroke=0)
        if BRAND.exists():
            self.c.drawImage(str(BRAND), MARGIN + 10, PAGE_H - MARGIN - 70, width=120, height=48, mask="auto", preserveAspectRatio=True)
        self.c.setFillColor(IVORY)
        self.c.setFont("Body", 9)
        self.c.drawString(MARGIN + 10, PAGE_H - MARGIN - 100, "SYSTEM TRAINING AND OPERATIONS GUIDE")
        self.c.setFont("Disp-Bold", 34)
        for i, line in enumerate(textwrap.wrap("REAP Formal Procurement Scorecard", 22)):
            self.c.drawString(MARGIN + 10, PAGE_H - MARGIN - 150 - i * 40, line)
        self.c.setFont("Body", 11)
        self.c.drawString(MARGIN + 10, PAGE_H / 2, "Instructor-led training  ·  Operating manual  ·  Quick reference")
        self.c.setStrokeColor(colors.HexColor("#C4A574"))
        self.c.setLineWidth(1.2)
        self.c.line(MARGIN + 10, PAGE_H / 2 - 16, MARGIN + 220, PAGE_H / 2 - 16)
        self.c.setFont("Body", 9)
        self.c.drawString(MARGIN + 10, MARGIN + 36, f"REAP Solutions  ·  {date.today().strftime('%B %Y')}")
        self.c.drawString(MARGIN + 10, MARGIN + 20, "Procurement Operations Ledger")

    def _kind_contents(self, p: Page):
        self._bg()
        y = PAGE_H - MARGIN
        self._section_label("Contents", y)
        y = self._title("Contents", y - 28, 24)
        y = self._rule(y)
        sections = [
            ("01", "Introduction — about the guide, concepts and journey"),
            ("02", "Access and navigation — public page, registration, sign-in, dashboard"),
            ("03", "Company management — directory, create, profile, history"),
            ("04", "Procurement assessments — TMPS, import, mapping, suppliers, save"),
            ("05", "Results and reporting — scorecard, categories, report, PDF, activity"),
            ("06", "Operations and support — settings, checklists, troubleshooting"),
        ]
        for code, label in sections:
            self.c.setFillColor(TEAL)
            self.c.setFont("Disp-Bold", 16)
            self.c.drawString(MARGIN, y, code)
            self.c.setFillColor(CHARCOAL)
            self.c.setFont("Body", 11)
            self.c.drawString(MARGIN + 50, y + 2, label)
            y -= 28
        y -= 10
        self.c.setFillColor(MUTED)
        self.c.setFont("Body", 8.5)
        self.c.drawString(MARGIN, y, "Keep the one-page quick reference beside this guide during live training.")

    def _kind_divider(self, p: Page):
        self.c.setFillColor(TEAL_DARK)
        self.c.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)
        self.c.setFillColor(IVORY)
        self.c.setFont("Body-Bold", 10)
        self.c.drawString(MARGIN, PAGE_H - MARGIN - 20, p.section)
        self.c.setFont("Disp-Bold", 36)
        for i, line in enumerate(textwrap.wrap(p.title, 24)):
            self.c.drawString(MARGIN, PAGE_H / 2 + 20 - i * 42, line)
        if p.purpose:
            self.c.setFont("Body", 11)
            self.c.drawString(MARGIN, PAGE_H / 2 - 50, p.purpose)

    def _kind_journey(self, p: Page):
        self._bg()
        y = PAGE_H - MARGIN
        self._section_label(p.section, y)
        y = self._title(p.title, y - 26)
        y = self._purpose(p.purpose, y)
        y = self._rule(y)
        cols = 4
        box_w = (PAGE_W - 2 * MARGIN - 30) / cols
        box_h = 58
        for i, label in enumerate(p.bullets):
            col = i % cols
            row = i // cols
            x = MARGIN + col * (box_w + 10)
            yy = y - row * (box_h + 16)
            self.c.setFillColor(WHITE)
            self.c.setStrokeColor(TEAL)
            self.c.setLineWidth(1)
            self.c.roundRect(x, yy - box_h, box_w, box_h, 4, fill=1, stroke=1)
            self.c.setFillColor(TEAL)
            self.c.setFont("Body-Bold", 9)
            self.c.drawString(x + 10, yy - 18, f"{i+1:02d}")
            self.c.setFillColor(CHARCOAL)
            self.c.setFont("Body", 8.5)
            for j, line in enumerate(textwrap.wrap(label, 28)):
                self.c.drawString(x + 10, yy - 36 - j * 11, line)

    def _kind_concept(self, p: Page):
        self._bg()
        y = PAGE_H - MARGIN
        self._section_label(p.section, y)
        y = self._title(p.title, y - 26)
        y = self._purpose(p.purpose, y)
        y = self._rule(y)
        left_w = PAGE_W - 2 * MARGIN
        if p.table:
            left_w = (PAGE_W - 2 * MARGIN) * 0.52
        for para in p.body:
            pr = Paragraph(para, styles["body"])
            w, h = pr.wrap(left_w, 200)
            pr.drawOn(self.c, MARGIN, y - h)
            y -= h + 8
        if p.bullets:
            y = self._steps_table(p.bullets, MARGIN, y, left_w)
        if p.table:
            data = [[Paragraph(f"<b>{c}</b>", styles["body_sm"]) if r == 0 else Paragraph(c, styles["body_sm"]) for c in row]
                    for r, row in enumerate(p.table)]
            tw = (PAGE_W - 2 * MARGIN) * 0.44
            t = Table(data, colWidths=[tw / len(p.table[0])] * len(p.table[0]))
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), TEAL_DARK),
                ("TEXTCOLOR", (0, 0), (-1, 0), WHITE),
                ("GRID", (0, 0), (-1, -1), 0.4, LINE),
                ("BACKGROUND", (0, 1), (-1, -1), WHITE),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]))
            # draw table on right
            rw = (PAGE_W - 2 * MARGIN) * 0.44
            ww, hh = t.wrap(rw, 400)
            t.drawOn(self.c, PAGE_W - MARGIN - rw, PAGE_H - MARGIN - 120 - hh)
        if p.note:
            self._panel("Important", p.note, MARGIN, max(y, MARGIN + 90), PAGE_W - 2 * MARGIN, fill=AMBER_PALE, title_color=BRASS)

    def _kind_checklist(self, p: Page):
        self._bg()
        y = PAGE_H - MARGIN
        self._section_label(p.section, y)
        y = self._title(p.title, y - 26)
        y = self._purpose(p.purpose, y)
        y = self._rule(y)
        for item in p.bullets:
            y = self._panel("Checklist", item, MARGIN, y, PAGE_W - 2 * MARGIN)

    def _kind_close(self, p: Page):
        self._bg()
        y = PAGE_H - MARGIN - 20
        self._section_label(p.section, y)
        y = self._title(p.title, y - 30, 26)
        y = self._rule(y - 4)
        for para in p.body:
            pr = Paragraph(para, styles["body"])
            w, h = pr.wrap(PAGE_W - 2 * MARGIN, 200)
            pr.drawOn(self.c, MARGIN, y - h)
            y -= h + 14
        self.c.setFillColor(TEAL)
        self.c.setFont("Disp-Bold", 12)
        self.c.drawString(MARGIN, MARGIN + 40, "REAP Solutions — Formal Procurement Scorecard")

    def _kind_ops(self, p: Page):
        self._bg()
        y = PAGE_H - MARGIN
        self._section_label(p.section, y)
        y = self._title(p.title, y - 24, 20)
        y = self._purpose(p.purpose, y)
        y = self._rule(y)

        left_w = (PAGE_W - 2 * MARGIN) * 0.42
        right_x = MARGIN + left_w + 16
        right_w = PAGE_W - MARGIN - right_x
        left_bottom = MARGIN + 24
        y_left = y
        if p.steps:
            self.c.setFillColor(TEAL_DARK)
            self.c.setFont("Body-Bold", 8)
            self.c.drawString(MARGIN, y_left, "OPERATING STEPS")
            y_left -= 14
            y_left = self._steps_table(p.steps, MARGIN, y_left, left_w)
        if p.checks:
            y_left = self._panel("What to check", "<br/>".join(f"• {c}" for c in p.checks), MARGIN, y_left, left_w)
        if p.mistake:
            y_left = self._panel("Common mistake", p.mistake, MARGIN, y_left, left_w, fill=AMBER_PALE, title_color=BRASS)
        if p.expected:
            y_left = self._panel("Expected result", p.expected, MARGIN, y_left, left_w)
        if p.note:
            y_left = self._panel("Note", p.note, MARGIN, y_left, left_w, fill=AMBER_PALE, title_color=BRASS)
        if p.table and not p.images:
            # rare: table on ops page without pushing images
            pass

        # images on right
        y_img = y
        imgs = [i for i in p.images if shot_path(i)]
        if not imgs:
            return
        if len(imgs) == 1:
            self._draw_image(imgs[0], right_x, y_img, right_w, y_img - left_bottom - 20, p.captions[0] if p.captions else "", p.fictional)
        elif len(imgs) == 2:
            h_each = (y_img - left_bottom - 30) / 2
            y_img = self._draw_image(imgs[0], right_x, y_img, right_w, h_each - 8, p.captions[0] if p.captions else "", p.fictional)
            self._draw_image(imgs[1], right_x, y_img - 6, right_w, h_each - 8, p.captions[1] if len(p.captions) > 1 else "", p.fictional)
        else:
            # context on top, details in a row below
            top_h = (y_img - left_bottom) * 0.58
            y_img = self._draw_image(imgs[0], right_x, y_img, right_w, top_h - 10, p.captions[0] if p.captions else "", p.fictional)
            detail_imgs = imgs[1:3]
            dw = (right_w - 10) / len(detail_imgs)
            for i, imn in enumerate(detail_imgs):
                cap = p.captions[i + 1] if len(p.captions) > i + 1 else ""
                self._draw_image(imn, right_x + i * (dw + 10), y_img - 4, dw, max(y_img - left_bottom - 10, 80), cap, False)


def write_quick_reference(path: Path):
    c = canvas.Canvas(str(path), pagesize=landscape(A4))
    c.setTitle("REAP Formal Procurement Scorecard — Quick Reference")
    c.setFillColor(IVORY)
    c.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)
    c.setFillColor(TEAL_DARK)
    c.rect(0, PAGE_H - 56, PAGE_W, 56, fill=1, stroke=0)
    c.setFillColor(IVORY)
    c.setFont("Disp-Bold", 18)
    c.drawString(MARGIN, PAGE_H - 36, "REAP Formal Procurement Scorecard — Quick Reference")
    c.setFont("Body", 9)
    c.drawRightString(PAGE_W - MARGIN, PAGE_H - 34, "One-page operating path")

    steps = [
        "Open or create a company",
        "Start a procurement assessment",
        "Select the assessment year",
        "Upload or paste supplier data",
        "Confirm column mapping",
        "Review supplier records",
        "Confirm the TMPS method",
        "Review the score preview",
        "Save the assessment",
        "Generate the report",
        "Download the PDF",
        "Review activity",
    ]
    cols = 3
    box_w = (PAGE_W - 2 * MARGIN - 24) / cols
    box_h = 42
    top = PAGE_H - 80
    for i, label in enumerate(steps):
        col = i % cols
        row = i // cols
        x = MARGIN + col * (box_w + 12)
        y = top - row * (box_h + 12)
        c.setFillColor(WHITE)
        c.setStrokeColor(TEAL)
        c.roundRect(x, y - box_h, box_w, box_h, 4, fill=1, stroke=1)
        c.setFillColor(TEAL)
        c.setFont("Body-Bold", 11)
        c.drawString(x + 12, y - 18, f"{i+1:02d}")
        c.setFillColor(CHARCOAL)
        c.setFont("Body", 9.5)
        c.drawString(x + 44, y - 18, label)

    # support panel
    y = 118
    c.setFillColor(TEAL_DARK)
    c.roundRect(MARGIN, 36, PAGE_W - 2 * MARGIN, y - 20, 5, fill=1, stroke=0)
    c.setFillColor(IVORY)
    c.setFont("Body-Bold", 10)
    c.drawString(MARGIN + 16, y - 8, "WHEN REQUESTING SUPPORT, PROVIDE:")
    items = ["Company name", "Assessment year", "Workbook filename", "Screenshot", "Exact error message", "Steps completed before the error"]
    c.setFont("Body", 9.5)
    for i, item in enumerate(items):
        c.drawString(MARGIN + 16 + (i % 3) * 240, y - 32 - (i // 3) * 18, f"•  {item}")
    c.showPage()
    c.save()


def write_docx(pages: list[Page], path: Path):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    for edge in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, edge, Cm(1.5))

    # footer
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("REAP Formal Procurement Scorecard · Operations Ledger")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x05, 0x2F, 0x32)
        return h

    for page in pages:
        if page.kind == "cover":
            add_heading("REAP Formal Procurement Scorecard", 0)
            doc.add_paragraph("System Training and Operations Guide")
            doc.add_paragraph("Instructor-led training · Operating manual · Quick reference")
            doc.add_paragraph(f"REAP Solutions · {date.today().strftime('%B %Y')}")
            doc.add_page_break()
            continue
        add_heading(page.title, 1)
        if page.section:
            sub = doc.add_paragraph(page.section)
            sub.runs[0].font.size = Pt(9)
            sub.runs[0].font.color.rgb = RGBColor(0x0B, 0x52, 0x59)
        if page.purpose:
            doc.add_paragraph(page.purpose)
        for para in page.body:
            doc.add_paragraph(para)
        if page.steps:
            add_heading("Operating steps", 2)
            for i, step in enumerate(page.steps, 1):
                doc.add_paragraph(f"{i}. {step}")
        if page.checks:
            add_heading("What to check", 2)
            for c in page.checks:
                doc.add_paragraph(c, style="List Bullet")
        if page.mistake:
            add_heading("Common mistake", 2)
            doc.add_paragraph(page.mistake)
        if page.expected:
            add_heading("Expected result", 2)
            doc.add_paragraph(page.expected)
        if page.bullets and page.kind in {"concept", "checklist", "journey"}:
            for b in page.bullets:
                doc.add_paragraph(b, style="List Bullet")
        if page.table:
            table = doc.add_table(rows=len(page.table), cols=len(page.table[0]))
            table.style = "Table Grid"
            for r, row in enumerate(page.table):
                for c, val in enumerate(row):
                    table.cell(r, c).text = val
        if page.fictional and page.images:
            note = doc.add_paragraph("FICTIONAL DEMONSTRATION DATA")
            note.runs[0].bold = True
            note.runs[0].font.color.rgb = RGBColor(0x8A, 0x5B, 0x00)
        for idx, img in enumerate(page.images):
            path_img = shot_path(img)
            if not path_img:
                continue
            # Width ~15cm keeps high effective DPI for 3840px sources
            doc.add_picture(str(path_img), width=Cm(15 if idx == 0 else 11))
            last = doc.paragraphs[-1]
            last.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cap = page.captions[idx] if idx < len(page.captions) else img
            cp = doc.add_paragraph(f"Figure — {cap}")
            cp.runs[0].font.size = Pt(8)
            cp.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            # accessible alt text
            try:
                doc.inline_shapes[-1]._inline.docPr.set("descr", cap)
            except Exception:
                pass
        if page.note:
            doc.add_paragraph(f"Note: {page.note}")
        doc.add_page_break()
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def write_markdown_files():
    MD_DIR.mkdir(parents=True, exist_ok=True)
    (MD_DIR / "REAP_Training_Facilitator_Notes.md").write_text(textwrap.dedent("""\
    # REAP Formal Procurement Scorecard — Facilitator Notes

    ## Recommended duration
    90–120 minutes for a full demonstration with guided practice, or 60 minutes for an experienced operator refresher.

    ## Practical agenda
    1. Platform introduction (10 min)
    2. Sign-in and navigation (10 min)
    3. Companies (10 min)
    4. Complete assessment demonstration (35–45 min)
    5. Score interpretation and reports (15 min)
    6. Guided user practice (15–20 min)
    7. Questions, feedback and closeout (10 min)

    ## Platform introduction
    - State that this session covers only the Formal Procurement Scorecard.
    - Show the public Scorecard page and the full journey map.
    - Emphasise: company → assessment year → suppliers → TMPS → score → save → report → activity.

    ## Sign-in and navigation
    - Demonstrate Create account only if attendees are new.
    - Sign in with the approved demonstration account.
    - Point out Dashboard, Companies, Activity, New Company, New Procurement Assessment, Settings, Need help?, Sign out.

    ## Companies
    - Create one fictional company or open the prepared demonstration company.
    - Stress required fields and the risk of duplicate company records.
    - Show the profile as the home for assessment history.

    ## Complete assessment demonstration
    What to demonstrate:
    - Assessment year
    - TMPS method (Calculated TMPS pad)
    - Workbook upload
    - Sheet selection and column mapping
    - Validation warnings and skipped rows
    - Supplier search and one inline edit
    - Live score preview
    - Save

    Suggested wording:
    - “Required mappings are Supplier name and Spend amount. Everything else improves category allocation.”
    - “TMPS is the denominator. If it is zero, the assessment cannot be saved.”
    - “Skipped rows are not ignored forever — they are evidence gaps.”

    ## Spreadsheet upload and column mapping
    - Use the fictional Thandeka workbook.
    - Pause on Found versus Needs mapping.
    - Show warnings for invalid or zero spend.

    ## Supplier correction and TMPS
    - Edit one supplier’s level or ownership flag and watch recognition change.
    - Reconfirm TMPS after any denominator discussion.

    ## Score interpretation and reports
    - Explain score out of 29 and procurement level.
    - Open category performance and recommendations.
    - View report → Download PDF.
    - Open Activity to show the audit trail.

    ## Guided user practice
    - Ask each trainee to repeat: open company → new assessment → map → save (or talk through if time-limited).
    - Questions to ask trainees:
      - Which fields are required when mapping?
      - What happens if TMPS is R0?
      - Why is recognised spend different from actual spend?
      - What must you send when requesting support?

    ## Common misunderstandings
    - Treating procurement level as the entity’s overall B-BBEE level.
    - Using a finance TMPS sheet as the supplier register.
    - Setting ownership flags without evidence.
    - Creating duplicate companies for spelling variants.

    ## Demonstration recovery steps
    - If import fails: reselect the Supplier Register sheet and remap required fields.
    - If save is blocked: check TMPS denominator and supplier spend.
    - If PDF fails: remain on the report page and retry Download PDF.
    - If the wrong company is open: return to Companies and open the correct profile.

    ## What not to claim
    - Do not claim the procurement result is a complete B-BBEE scorecard.
    - Do not present illustrative recognition percentages as legal advice.
    - Do not promise points the evidence cannot support.

    ## Training closeout checklist
    - Guide and quick reference distributed
    - Support contact confirmed
    - Unresolved questions logged
    - Defects separated from feature requests
    - Demonstration data labelled fictional
    """), encoding="utf-8")

    (MD_DIR / "REAP_Training_Checklist.md").write_text(textwrap.dedent("""\
    # REAP Formal Procurement Scorecard — Training Checklist

    ## Before training
    - [ ] Application tested
    - [ ] Login tested
    - [ ] Demo account available
    - [ ] Demo company available
    - [ ] Supplier workbook ready
    - [ ] PDF export tested
    - [ ] Browser zoom set to 100%
    - [ ] Charger available
    - [ ] Internet fallback considered
    - [ ] Training guide available offline
    - [ ] Quick-reference sheet available
    - [ ] Confidential data removed

    ## During training
    - [ ] Explain the full workflow
    - [ ] Create one company
    - [ ] Complete one assessment
    - [ ] Upload supplier data
    - [ ] Map the columns
    - [ ] Review import validation
    - [ ] Edit one supplier
    - [ ] Confirm TMPS
    - [ ] Review the score
    - [ ] Save the assessment
    - [ ] Generate a report
    - [ ] Download the PDF
    - [ ] Allow users to repeat the workflow
    - [ ] Record questions
    - [ ] Record requested changes

    ## After training
    - [ ] Send the guide
    - [ ] Send the quick reference
    - [ ] Confirm support contacts
    - [ ] Record unresolved questions
    - [ ] Separate defects from new feature requests
    - [ ] Prepare a change scope where necessary
    """), encoding="utf-8")

    (MD_DIR / "REAP_Troubleshooting_Guide.md").write_text(textwrap.dedent("""\
    # REAP Formal Procurement Scorecard — Troubleshooting Guide

    ## Cannot sign in
    **Problem:** The sign-in form rejects the credentials or returns to the login screen.  
    **Possible cause:** Wrong email/password, unconfirmed account, or expired session.  
    **How to resolve it:** Confirm the email address, use Forgot password, wait for the reset link, then set a new password.  
    **Information to provide to support:** Email used, exact error message, time of attempt.

    ## Password reset problem
    **Problem:** Reset link is invalid or expired.  
    **Possible cause:** Link already used or older than the allowed window.  
    **How to resolve it:** Request a new reset from Forgot password and open the newest email.  
    **Information to provide to support:** Email address and screenshot of the expired-link message.

    ## Company not appearing
    **Problem:** Expected company is missing from Companies.  
    **Possible cause:** Different signed-in user, company not yet created, or wrong environment.  
    **How to resolve it:** Confirm the account, refresh Companies, create the company if needed.  
    **Information to provide to support:** Company name, account email, screenshot of the directory.

    ## Duplicate company
    **Problem:** Two records exist for one organisation.  
    **Possible cause:** Spelling variants created on different days.  
    **How to resolve it:** Open both profiles, keep assessments on the correct record, stop using the duplicate, ask support about clean-up if required.  
    **Information to provide to support:** Both company names and which record should remain.

    ## Workbook rejected
    **Problem:** Upload fails.  
    **Possible cause:** Wrong file type or file too large.  
    **How to resolve it:** Use .xlsx or .xls and keep the file within the allowed size.  
    **Information to provide to support:** Filename, size, exact error text.

    ## Wrong file type
    **Problem:** System says the file type is unsupported.  
    **Possible cause:** CSV, PDF or macro-enabled format uploaded.  
    **How to resolve it:** Save as .xlsx and upload again.  
    **Information to provide to support:** Original format and converted filename.

    ## Missing required columns
    **Problem:** Mapping shows Needs mapping for required fields.  
    **Possible cause:** Header names not recognised or wrong sheet selected.  
    **How to resolve it:** Select the supplier register sheet and map Supplier name and Spend amount manually.  
    **Information to provide to support:** Screenshot of Column mapping and the header row.

    ## Supplier rows not importing
    **Problem:** No suppliers loaded.  
    **Possible cause:** Header-only sheet, wrong sheet, or non-numeric spend.  
    **How to resolve it:** Confirm data rows exist under the header; remap spend; apply suppliers again.  
    **Information to provide to support:** Sheet name, rows read, warnings text.

    ## Invalid supplier values
    **Problem:** Rows skipped for invalid or zero spend.  
    **Possible cause:** Text in spend cells, blanks, or zero amounts.  
    **How to resolve it:** Correct the workbook or edit/add the supplier manually after import.  
    **Information to provide to support:** Warning lines and the affected supplier names.

    ## Unexpected procurement score
    **Problem:** Points differ from expectation.  
    **Possible cause:** Wrong TMPS method/value, incorrect levels, or ownership flags.  
    **How to resolve it:** Recheck TMPS, recognition levels, ownership evidence and live preview before saving.  
    **Information to provide to support:** Assessment year, TMPS value, screenshot of score preview.

    ## Wrong B-BBEE level
    **Problem:** Supplier level looks incorrect.  
    **Possible cause:** Bad import mapping or outdated certificate data.  
    **How to resolve it:** Edit the supplier row and set the evidenced level.  
    **Information to provide to support:** Supplier name and certificate evidence reference.

    ## Wrong ownership information
    **Problem:** Black Owned / Black Women Owned / Designated Group flags are wrong.  
    **Possible cause:** Incorrect import values or manual error.  
    **How to resolve it:** Clear or set flags only where evidence exists, then recalculate.  
    **Information to provide to support:** Supplier name and which flag is wrong.

    ## Wrong TMPS method
    **Problem:** Denominator does not match the agreed schedule.  
    **Possible cause:** Calculated pad selected instead of supplier-spend method, or vice versa.  
    **How to resolve it:** Change the TMPS method, re-enter pad values if needed, review live summary.  
    **Information to provide to support:** Intended method and screenshot of the TMPS panel.

    ## Assessment not saving
    **Problem:** Save remains disabled or errors.  
    **Possible cause:** Zero TMPS, no suppliers, or validation errors.  
    **How to resolve it:** Fix the denominator, add at least one positive-spend supplier, clear Before you can save messages.  
    **Information to provide to support:** Exact blocking message.

    ## Assessment cannot be reopened
    **Problem:** Edit or open fails.  
    **Possible cause:** Wrong company, deleted assessment, or session expired.  
    **How to resolve it:** Return via company history, refresh, sign in again if needed.  
    **Information to provide to support:** Company name, assessment year, URL if visible.

    ## PDF not downloading
    **Problem:** Download PDF does not produce a file.  
    **Possible cause:** Browser blocking downloads or report page not ready.  
    **How to resolve it:** Stay on the report page, retry Download PDF, check browser download permissions.  
    **Information to provide to support:** Browser name, screenshot, assessment year.

    ## Session expired
    **Problem:** User is returned to sign-in.  
    **Possible cause:** Idle timeout.  
    **How to resolve it:** Sign in again and reopen the company/assessment from history.  
    **Information to provide to support:** Time idle and last completed step.

    ## Page not loading
    **Problem:** Blank or endless loading state.  
    **Possible cause:** Network interruption or temporary application error.  
    **How to resolve it:** Refresh once, try another browser, confirm network.  
    **Information to provide to support:** URL path, browser, screenshot.

    ## User needs support
    **Problem:** Issue remains after the checks above.  
    **Possible cause:** Defect or environment-specific fault.  
    **How to resolve it:** Escalate with the support pack.  
    **Information to provide to support:** Company name; assessment year; workbook filename; screenshot; exact error message; steps completed before the error.
    """), encoding="utf-8")


def write_screenshot_index(used: dict[str, list[int]]):
    ASSET_OUT.mkdir(parents=True, exist_ok=True)
    lines = [
        "# REAP Formal Procurement Scorecard — Screenshot Index",
        "",
        f"Capture date: {date.today().isoformat()}",
        "Capture method: Chromium via Playwright, 1920×1080 CSS pixels, device scale factor 2, lossless PNG.",
        "Demonstration company: Thandeka Industrial Holdings (Pty) Ltd (fictional).",
        "",
        "| File | Dimensions | Size (KB) | Route | Purpose | Type | Guide pages |",
        "|---|---|---|---|---|---|---|",
    ]
    manifest = SHOT / "_capture-manifest.json"
    meta = {}
    if manifest.exists():
        import json
        meta = {s["file"]: s for s in __import__("json").loads(manifest.read_text())["shots"]}
    for path in sorted(SHOT.glob("*.png")):
        if path.name.startswith("_"):
            continue
        im = Image.open(path)
        m = meta.get(path.name, {})
        pages = ",".join(str(p) for p in used.get(path.name, [])) or "—"
        purpose = str(m.get("purpose", "Training capture")).replace(
            "Live procurement score preview", "Live score preview"
        )
        lines.append(
            f"| {path.name} | {im.size[0]}×{im.size[1]} | {path.stat().st_size//1024} | "
            f"{m.get('route','—')} | {purpose} | {m.get('kind','—')} | {pages} |"
        )
    # omissions
    lines += [
        "",
        "## Coverage notes",
        "",
        "- Company directory search: not present in the current user-facing application; browsing and careful record selection are documented instead.",
        "- Manual fixed TMPS amount: not exposed in the current UI; Calculated TMPS and Use supplier spend as TMPS are documented.",
        "- All screenshots were captured from the current Formal Procurement Scorecard application.",
        "",
    ]
    INDEX_MD.write_text("\n".join(lines), encoding="utf-8")


def render_qa(pdf_path: Path, out_dir: Path):
    out_dir.mkdir(parents=True, exist_ok=True)
    for old in out_dir.glob("*.png"):
        old.unlink()
    doc = fitz.open(pdf_path)
    for i, page in enumerate(doc, 1):
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        pix.save(out_dir / f"page-{i:02d}.png")
    return len(doc)


def effective_dpi_checks(pdf_path: Path) -> list[str]:
    issues = []
    doc = fitz.open(pdf_path)
    for i, page in enumerate(doc, 1):
        for img in page.get_images(full=True):
            xref = img[0]
            try:
                info = doc.extract_image(xref)
            except Exception:
                continue
            w, h = info.get("width", 0), info.get("height", 0)
            if w < 400 or h < 40:
                # likely icon; skip
                continue
            # Approximate display size via first image rect
            rects = page.get_image_rects(xref)
            if not rects:
                continue
            rect = rects[0]
            disp_in_w = max(rect.width / 72, 0.01)
            dpi = w / disp_in_w
            if dpi < 140 and w >= 800:
                issues.append(f"page {i}: image ~{dpi:.0f} DPI effective ({w}px over {disp_in_w:.2f}in)")
    return issues


def prohibited_scan(paths: list[Path]) -> list[str]:
    hits = []
    for path in paths:
        if not path.exists():
            continue
        if path.suffix.lower() == ".pdf":
            doc = fitz.open(path)
            text = "\n".join(page.get_text() for page in doc)
        else:
            text = path.read_text(encoding="utf-8", errors="ignore")
        for m in PROHIBITED.finditer(text):
            hits.append(f"{path.name}: '{m.group(0)}'")
    return hits


def build_zip():
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    staging = ROOT / "output" / "_zip_staging" / "REAP_Scorecard_Training_Package"
    if staging.parent.exists():
        shutil.rmtree(staging.parent)
    staging.mkdir(parents=True)
    mapping = [
        (MAIN_PDF, staging / MAIN_PDF.name),
        (MAIN_DOCX, staging / MAIN_DOCX.name),
        (QUICK_PDF, staging / QUICK_PDF.name),
        (MD_DIR / "REAP_Training_Facilitator_Notes.md", staging / "REAP_Training_Facilitator_Notes.md"),
        (MD_DIR / "REAP_Training_Checklist.md", staging / "REAP_Training_Checklist.md"),
        (MD_DIR / "REAP_Troubleshooting_Guide.md", staging / "REAP_Troubleshooting_Guide.md"),
    ]
    assets = staging / "reap-training-assets"
    assets.mkdir()
    mapping.append((INDEX_MD, assets / "screenshot-index.md"))
    for src, dst in mapping:
        shutil.copy2(src, dst)
    with zipfile.ZipFile(ZIP_PATH, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in staging.rglob("*"):
            if path.is_file():
                zf.write(path, path.relative_to(staging.parent).as_posix())
    shutil.rmtree(staging.parent)
    # verify members
    with zipfile.ZipFile(ZIP_PATH) as zf:
        names = sorted(zf.namelist())
    return names


def main():
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    pages = build_pages()
    print(f"Building guide with {len(pages)} pages…")
    guide = GuidePDF(pages, MAIN_PDF)
    guide.draw()
    print(f"Wrote {MAIN_PDF}")

    write_quick_reference(QUICK_PDF)
    print(f"Wrote {QUICK_PDF}")

    write_docx(pages, MAIN_DOCX)
    print(f"Wrote {MAIN_DOCX}")

    write_markdown_files()
    write_screenshot_index(guide.used_images)
    print(f"Wrote markdown + {INDEX_MD}")

    main_pages = render_qa(MAIN_PDF, QA_MAIN)
    quick_pages = render_qa(QUICK_PDF, QA_QUICK)
    dpi_issues = effective_dpi_checks(MAIN_PDF)
    hits = prohibited_scan([
        MAIN_PDF, QUICK_PDF, MAIN_DOCX,
        MD_DIR / "REAP_Training_Facilitator_Notes.md",
        MD_DIR / "REAP_Training_Checklist.md",
        MD_DIR / "REAP_Troubleshooting_Guide.md",
        INDEX_MD,
    ])
    names = build_zip()
    expected = {
        "REAP_Scorecard_Training_Package/REAP_Scorecard_System_Training_Guide.pdf",
        "REAP_Scorecard_Training_Package/REAP_Scorecard_System_Training_Guide.docx",
        "REAP_Scorecard_Training_Package/REAP_Scorecard_Quick_Reference.pdf",
        "REAP_Scorecard_Training_Package/REAP_Training_Facilitator_Notes.md",
        "REAP_Scorecard_Training_Package/REAP_Training_Checklist.md",
        "REAP_Scorecard_Training_Package/REAP_Troubleshooting_Guide.md",
        "REAP_Scorecard_Training_Package/reap-training-assets/screenshot-index.md",
    }

    print("\n===== PREFLIGHT =====")
    ok = True
    def check(label, cond, detail=""):
        nonlocal ok
        status = "PASS" if cond else "FAIL"
        if not cond:
            ok = False
        print(f"{status}: {label} {detail}")

    check("Main PDF page count 38–46", 38 <= main_pages <= 46, f"({main_pages})")
    check("Quick reference exactly 1 page", quick_pages == 1, f"({quick_pages})")
    check("No prohibited terms", not hits, f"({hits[:5]})" if hits else "")
    check("ZIP members exact", set(names) == expected, f"({names})")
    check("Screenshot DPI floor", len(dpi_issues) == 0, f"({dpi_issues[:3]})" if dpi_issues else "")
    check("ZIP exists", ZIP_PATH.exists(), f"({ZIP_PATH.stat().st_size//1024} KB)")
    print("PREFLIGHT", "PASS" if ok else "FAIL")
    print(f"Main pages rendered to {QA_MAIN}")
    print(f"Quick pages rendered to {QA_QUICK}")


if __name__ == "__main__":
    main()
